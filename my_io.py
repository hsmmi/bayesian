import os
from random import randint
import numpy as np
import pandas as pd

def build_testcase(N,p,q,Range):
    with open(os.path.join(os.path.dirname(__file__),"HW1DS.txt"),'w') as f:
        f.write(f'p in norm of all deviations is: {p}\n')

    with open(os.path.join(os.path.dirname(__file__),"HW1DS.txt"),'a') as f:
        f.write(f'q in deviation of each part is: {q}\n')

    tmp = [None] * N
    for i in range(N):
        tmp[i] = randint(1, Range)-Range/2
    tmp.sort()

    with open(os.path.join(os.path.dirname(__file__),"HW1DS.txt"),'a') as f:
        f.write(f'Points are: {tmp}\n')

def read_testcase(file, printer = 0):
    with open(os.path.join(os.path.dirname(__file__),file),'r') as f:
        p = int(f.readline().split(':')[1])
        q = int(f.readline().split(':')[1])
        points = list(map(float, f.readline().split(':')[1][2:-2].split(',')))
    points.sort()

    if(printer):
        print(f'p in norm of all deviations is:\n{p}\n')
        print(f'q in deviation of each part is:\n{q}\n')
        print(f'Points are:\n{points}\n')
    return p, q, points

def read_dataset(file, atr):
    if (type(atr) == int):
        with open(os.path.join(os.path.dirname(__file__),file),'r') as f:
            return list(map(lambda x: float(x.split(',')[atr]), f.read().splitlines()))
    else:
        with open(os.path.join(os.path.dirname(__file__),file),'r') as f:
            return list(map(lambda x: [float(i) for i in (x.split(',')[atr[0]:atr[1]])], f.read().splitlines()))

def read_dataset_with_pandas(file, atr= None):
    colName = pd.read_csv(os.path.join(os.path.dirname(__file__),file),nrows=0).columns
    if (type(atr) == int):
        colName = [colName[atr]]
    elif(atr != None):
        colName = colName[atr[0]:atr[1]]
    data = pd.read_csv(os.path.join(os.path.dirname(__file__),file),usecols=colName)

    return colName, data

def dataframe_to_docx_table(header,data,file,doc=None,save=1):
    """
    Read header and data
    If you gave if doc it add header and data to it and return it
    If you gave it save=0 it will not be save doc
    Return doc include header and data
    """
    import docx
    if(doc == None):
        doc = docx.Document()
    doc.add_heading(header, 1)

    table = doc.add_table(rows=len(data.index)+1, cols=len(data.columns)+1)

    for j in range(len(data.columns)):
        table.cell(0,j+1).text = f'{data.columns[j]}'

    for i in range(len(data.index)):
        table.cell(i+1,0).text = f'{data.index[i]}'
        for j in range(len(data.columns)):
            table.cell(i+1,j+1).text = f'{data.iat[i,j]}'
    table.style = 'Table Grid'
    if(save):
        doc.save(file)
    return doc

def string_to_dataframe(string):
    from io import StringIO
    data = StringIO(string)
    return pd.read_csv(data)

def generate_dataset(
    file: str, mean_ds: np.ndarray, cov_ds: np.ndarray,
        sampels_size: int) -> pd.core.frame.DataFrame:
    dataset = pd.DataFrame(data={'X1': [], 'X2': [], 'Y': []})
    for i in range(mean_ds.shape[0]):
        x1, x2 = np.random.multivariate_normal(
            mean_ds[i], cov_ds[i], sampels_size).T
        temp = pd.DataFrame(data={'X1': x1, 'X2': x2, 'Y': [i]*sampels_size})
        dataset = pd.concat([dataset, temp], axis=0)
    dataset.to_csv(file, index=False)
    return dataset
